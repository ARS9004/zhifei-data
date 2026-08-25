#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智飞投研 · 云端 v3.8 (2026-08-25)
- 后加载-智能接续：新会话启动时，取最后5轮完整对话 + 最后25轮摘要
- 每轮写入：前端发指令，模型 executeQuerySql 写入 RDS
- 依赖百炼 session 记忆，启动恢复后正常对话由百炼接管
"""

import os
import re
import json
import time
import uuid
import logging
import io
import threading
from datetime import datetime
from typing import List, Dict

import streamlit as st
import dashscope
from dashscope import Application
import pytz
from http import HTTPStatus
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

load_dotenv()

_CHINESE_CHAR_RE = re.compile(r'[\u4e00-\u9fff]')

def get_secret_or_env(key, secrets_key=None, default=None):
    if secrets_key:
        parts = secrets_key.split('.')
        try:
            value = st.secrets
            for p in parts:
                value = value[p]
            if value:
                return value
        except Exception:
            pass
    return os.getenv(key, default)

DASHSCOPE_API_KEY = get_secret_or_env("DASHSCOPE_API_KEY", "dashscope.api_key")
if not DASHSCOPE_API_KEY:
    raise RuntimeError("请配置 DASHSCOPE_API_KEY")

MODEL_NAME = get_secret_or_env("MODEL_NAME", "model.name", "deepseek-v4-pro")
BEIJING_TZ = pytz.timezone('Asia/Shanghai')

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def now_ts_display():
    return datetime.now(BEIJING_TZ).strftime("%Y-%m-%d %H:%M:%S")

def is_one_click_scheme(scheme_name):
    return scheme_name in ["盘前快速分析", "产业链扫描", "简报", "周报"]

def get_scheme_prompt(scheme_name):
    return f"请执行【{scheme_name}】分析方案"

# ================= 熔断机制 =================
_FAIL_LOCK = threading.Lock()
_FAIL_COUNTER = {"network": 0, "api": 0, "model": 0}
_MODEL_HEALTHY = True

def reset_health_status():
    global _FAIL_COUNTER, _MODEL_HEALTHY
    with _FAIL_LOCK:
        _FAIL_COUNTER = {"network": 0, "api": 0, "model": 0}
        _MODEL_HEALTHY = True

def mark_failure(error_type: str):
    global _FAIL_COUNTER, _MODEL_HEALTHY
    if error_type not in _FAIL_COUNTER:
        return
    with _FAIL_LOCK:
        _FAIL_COUNTER[error_type] += 1
        if error_type == "model" and _FAIL_COUNTER["model"] >= 3:
            _MODEL_HEALTHY = False
            logger.warning("模型服务熔断")

def is_model_healthy() -> bool:
    with _FAIL_LOCK:
        return _MODEL_HEALTHY

def _classify_error(e: Exception) -> str:
    err_str = str(e).lower()
    if not err_str:
        return "network"
    if any(kw in err_str for kw in ["timeout", "connection", "network", "resolve", "refused"]):
        return "network"
    if any(kw in err_str for kw in ["rate", "quota", "throttle", "limit", "429"]):
        return "api"
    return "model"

# ================= 百炼调用 =================
def _clean_for_api(raw_msgs: List[Dict]) -> List[Dict]:
    valid = []
    for m in raw_msgs:
        role = m.get("role", "")
        content = m.get("content")
        if role not in ("user", "assistant", "system"):
            continue
        if content is None or str(content).strip() == "":
            continue
        valid.append({"role": role, "content": str(content).strip()})
    if not valid:
        return []
    return valid

def _extract_text_from_response(resp) -> str:
    output = resp.output
    text = getattr(output, 'text', None)
    if text and isinstance(text, str) and text.strip():
        return text
    choices = getattr(output, 'choices', None)
    if choices and isinstance(choices, list) and len(choices) > 0:
        msg = getattr(choices[0], 'message', None)
        if msg:
            content = getattr(msg, 'content', None)
            if content and isinstance(content, str) and content.strip():
                return content
        assistant = getattr(choices[0], 'assistant', None)
        if assistant:
            if isinstance(assistant, dict):
                content = assistant.get('content') or assistant.get('text')
                if content and isinstance(content, str) and content.strip():
                    return content
            elif isinstance(assistant, str) and assistant.strip():
                return assistant
    if isinstance(output, str) and output.strip():
        return output
    if isinstance(output, dict):
        text = output.get('text')
        if not (text and isinstance(text, str) and text.strip()):
            text = output.get('content')
        if text and isinstance(text, str) and text.strip():
            return text
    return ""

def build_ctx_restore_instruction() -> str:
    """后加载-智能接续：新会话启动时执行一次"""
    return """新会话启动，接上之前的内容：

从 chat_memory 最后取两段数据：
- 最后5轮完整对话（做无缝对接）
- 最后25轮对话数据（用来写摘要）

用那25轮写一篇500字左右的摘要，把最近聊了什么主线、涉及的标的和板块、关键结论都串起来。

回复的时候先贴摘要，再贴最后5轮完整对话，然后正常回我消息。"""

def build_write_instruction(user_content: str, assistant_content: str) -> str:
    """构造对话写入指令（每轮结束后执行）"""
    user_escaped = json.dumps(user_content, ensure_ascii=False)
    assistant_escaped = json.dumps(assistant_content, ensure_ascii=False)
    return f"""[系统] 请将本轮对话原文写入 rds chat_memory：

使用 executeQuerySql 工具操作 rds chat_memory 表，追加新增两条 jsonl 数据：

INSERT INTO chat_memory (jsonl) VALUES ('{json.dumps({"role": "user", "content": user_escaped}, ensure_ascii=False)}');
INSERT INTO chat_memory (jsonl) VALUES ('{json.dumps({"role": "assistant", "content": assistant_escaped}, ensure_ascii=False)}');

要求：
- 原文逐字写入，不截断
- 追加新增，不影响已有数据"""

def call_bailian(messages: List[Dict]) -> str:
    if not is_model_healthy():
        raise RuntimeError("服务暂时不可用")
    dashscope.api_key = DASHSCOPE_API_KEY
    cleaned = _clean_for_api(messages)
    if not cleaned:
        raise RuntimeError("上下文中没有有效的用户消息")
    BAILIAN_APP_ID = "45db2f797bfd49229f757b04ed13ac92"
    retries, delay = 3, 2
    for attempt in range(retries):
        try:
            resp = Application.call(
                app_id=BAILIAN_APP_ID,
                messages=cleaned,
                stream=False
            )
            if resp.status_code == HTTPStatus.OK:
                full_text = _extract_text_from_response(resp)
                if not full_text or not full_text.strip():
                    raise RuntimeError("模型返回内容为空")
                reset_health_status()
                return full_text
            else:
                raise RuntimeError(f"API Error: {resp.code} - {resp.message}")
        except Exception as e:
            err_type = _classify_error(e)
            if attempt == retries - 1:
                mark_failure(err_type)
                raise RuntimeError(f"模型连续{retries}次调用失败: {e}")
            logger.warning(f"call_bailian 重试 {attempt+1}/{retries}: {e}")
            time.sleep(delay)
            delay *= 2
    raise RuntimeError("未知错误")

# ================= 导出 =================
def export_docx(messages):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '仿宋')

    title = doc.add_heading('', level=1)
    title_run = title.add_run("智飞投研分析方案")
    title_run.font.name = '仿宋'
    title_run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '仿宋')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    for m in messages:
        role = "用户" if m.get("role") == "user" else "智飞"
        content = m.get("content") or ""
        ts = m.get("timestamp", "")[:16]
        prefix = f"[{ts}] " if ts else ""
        para = doc.add_paragraph()
        run = para.add_run(f"{prefix}{role}：{content}")
        run.font.name = '仿宋'
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(14)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= 会话管理 =================
def get_history_sessions():
    """从 RDS chat_memory 读取所有非删除会话列表"""
    try:
        import pymysql
        conn = pymysql.connect(
            host=get_secret_or_env("RDS_HOST", "rds.host", "rm-2zeli1or40iqt7vq66o.mysql.rds.aliyuncs.com"),
            port=int(get_secret_or_env("RDS_PORT", "rds.port", "3306")),
            user=get_secret_or_env("RDS_USER", "rds.user", "zhuanz1"),
            password=get_secret_or_env("RDS_PASSWORD", "rds.password", ""),
            database="stock_db",
            charset="utf8mb4",
            connect_timeout=5
        )
        cursor = conn.cursor()
        cursor.execute("""
            SELECT cm.session_id, COUNT(*) as rounds, MAX(cm.created_at) as last_time,
                   (SELECT cm2.content FROM chat_memory cm2 
                    WHERE cm2.session_id = cm.session_id AND cm2.role = 'user' 
                    ORDER BY cm2.id ASC LIMIT 1) as first_msg
            FROM chat_memory cm
            WHERE cm.session_id NOT IN (SELECT session_id FROM deleted_sessions)
            GROUP BY cm.session_id
            ORDER BY last_time DESC
        """)
        rows = cursor.fetchall()
        conn.close()
        sessions = []
        for row in rows:
            sid, rounds, last_time, first_msg = row
            title = (first_msg or "")[:30] if first_msg else "（无标题）"
            sessions.append({
                "session_id": sid,
                "title": title,
                "rounds": rounds,
                "last_time": str(last_time) if last_time else ""
            })
        return sessions
    except Exception as e:
        logger.warning(f"get_history_sessions 失败: {e}")
        return []

def load_session_messages(session_id: str) -> List[Dict]:
    """从 RDS chat_memory 加载指定会话的所有消息"""
    try:
        import pymysql
        conn = pymysql.connect(
            host=get_secret_or_env("RDS_HOST", "rds.host", "rm-2zeli1or40iqt7vq66o.mysql.rds.aliyuncs.com"),
            port=int(get_secret_or_env("RDS_PORT", "rds.port", "3306")),
            user=get_secret_or_env("RDS_USER", "rds.user", "zhuanz1"),
            password=get_secret_or_env("RDS_PASSWORD", "rds.password", ""),
            database="stock_db",
            charset="utf8mb4",
            connect_timeout=5
        )
        cursor = conn.cursor()
        cursor.execute(
            "SELECT role, content, created_at FROM chat_memory WHERE session_id=%s ORDER BY id ASC",
            (session_id,)
        )
        rows = cursor.fetchall()
        conn.close()
        messages = []
        for row in rows:
            role, content, created_at = row
            messages.append({
                "role": role,
                "content": content or "",
                "timestamp": str(created_at) if created_at else ""
            })
        return messages
    except Exception as e:
        logger.warning(f"load_session_messages 失败: {e}")
        return []

def delete_session(session_id: str):
    """软删除会话：写入 deleted_sessions 表"""
    try:
        import pymysql
        conn = pymysql.connect(
            host=get_secret_or_env("RDS_HOST", "rds.host", "rm-2zeli1or40iqt7vq66o.mysql.rds.aliyuncs.com"),
            port=int(get_secret_or_env("RDS_PORT", "rds.port", "3306")),
            user=get_secret_or_env("RDS_USER", "rds.user", "zhuanz1"),
            password=get_secret_or_env("RDS_PASSWORD", "rds.password", ""),
            database="stock_db",
            charset="utf8mb4",
            connect_timeout=5
        )
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO deleted_sessions(session_id, deleted_at) VALUES(%s, NOW())",
            (session_id,)
        )
        conn.commit()
        conn.close()
        logger.info(f"会话已删除: {session_id}")
        return True
    except Exception as e:
        logger.warning(f"delete_session 失败: {e}")
        return False

# ================= 初始化 =================
def init_session():
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "display_messages" not in st.session_state:
        st.session_state.display_messages = []
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if "render_offset" not in st.session_state:
        st.session_state.render_offset = 0
    if "model_name" not in st.session_state:
        st.session_state.model_name = MODEL_NAME
    if "scheme" not in st.session_state:
        st.session_state.scheme = "日常对话"
    if "generating" not in st.session_state:
        st.session_state.generating = False
    if "stop" not in st.session_state:
        st.session_state.stop = False
    if "total_tokens_used" not in st.session_state:
        st.session_state.total_tokens_used = 0
    if "history_loaded" not in st.session_state:
        st.session_state.history_loaded = False
    if "ctx_restoring" not in st.session_state:
        st.session_state.ctx_restoring = False

# ================= 上下文恢复（启动时执行一次）=================
def trigger_ctx_restore():
    if st.session_state.history_loaded:
        return
    if st.session_state.messages:
        st.session_state.history_loaded = True
        return

    st.session_state.ctx_restoring = True
    st.session_state.generating = True

    ctx = [
        {"role": "system", "content": "你是智飞投研助手。"},
        {"role": "user", "content": build_ctx_restore_instruction()}
    ]

    try:
        with st.spinner("正在恢复上下文..."):
            reply = call_bailian(ctx)

        restore_msg = {
            "role": "assistant",
            "content": reply,
            "timestamp": now_ts_display()
        }
        st.session_state.messages.append(restore_msg)
        st.session_state.display_messages.append(restore_msg)

    except Exception as e:
        logger.warning(f"上下文恢复失败: {e}")

    st.session_state.history_loaded = True
    st.session_state.ctx_restoring = False
    st.session_state.generating = False

# ================= UI =================
st.set_page_config(page_title="智飞投研·云端", layout="centered")
st.title("智飞投研")

init_session()

if not st.session_state.history_loaded:
    trigger_ctx_restore()
    st.rerun()

total_rounds = len([m for m in st.session_state.messages if m["role"] == "user"])
st.caption(f"{total_rounds} 轮对话")

chat_container = st.container()
with chat_container:
    MAX_RENDER_MSGS = 60
    total_msgs = len(st.session_state.display_messages)
    render_count = (st.session_state.get("render_offset", 0) + 1) * MAX_RENDER_MSGS
    render_start = max(0, total_msgs - render_count)
    render_msgs = st.session_state.display_messages[render_start:] if st.session_state.display_messages else []

    if render_start > 0:
        if st.button("加载更早的对话", use_container_width=True):
            st.session_state.render_offset = st.session_state.get("render_offset", 0) + 1
            st.rerun()

    for m in render_msgs:
        with st.chat_message(m.get("role", "user")):
            content = str(m.get("content") or "")
            if content.startswith("❌"):
                st.error(content)
            else:
                st.markdown(content)

_pending_prompt = st.session_state.pop("_pending_prompt", None)
user_input = st.chat_input("输入消息...", disabled=st.session_state.generating)
if _pending_prompt and not st.session_state.generating:
    user_input = _pending_prompt

if user_input and not st.session_state.generating:
    st.session_state.generating = True
    user_msg = {
        "role": "user",
        "content": user_input,
        "timestamp": now_ts_display()
    }
    st.session_state.messages.append(user_msg)
    st.session_state.display_messages.append(user_msg)
    st.rerun()

if st.session_state.generating and st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    if not st.session_state.ctx_restoring:
        user_msg = st.session_state.messages[-1]
        user_content = str(user_msg.get("content") or "")

        ctx = [{"role": "user", "content": user_content}]

        try:
            with st.spinner("思考中..."):
                reply = call_bailian(ctx)

            assistant_msg = {
                "role": "assistant",
                "content": reply,
                "timestamp": now_ts_display()
            }
            st.session_state.messages.append(assistant_msg)
            st.session_state.display_messages.append(assistant_msg)

            write_instruction = build_write_instruction(user_content, reply)
            write_ctx = [
                {"role": "system", "content": "你是智飞投研助手，执行 SQL 写入操作。"},
                {"role": "user", "content": write_instruction}
            ]
            try:
                call_bailian(write_ctx)
            except Exception as e:
                logger.warning(f"写入RDS失败: {e}")

        except Exception as e:
            st.error(f"错误: {e}")
            err_msg = {
                "role": "assistant",
                "content": f"调用失败: {e}",
                "timestamp": now_ts_display()
            }
            st.session_state.messages.append(err_msg)
            st.session_state.display_messages.append(err_msg)

    st.session_state.generating = False
    st.rerun()

st.divider()

col1, col2, col3 = st.columns(3)
with col1:
    if st.button("新建会话", use_container_width=True, disabled=st.session_state.generating):
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.messages = []
        st.session_state.display_messages = []
        st.session_state.render_offset = 0
        st.session_state.history_loaded = False
        st.rerun()

with col2:
    if st.button("重新生成", use_container_width=True, disabled=st.session_state.generating):
        if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
            st.session_state.messages.pop()
            st.session_state.display_messages.pop()
            st.session_state.generating = True
            st.rerun()

with col3:
    st.download_button(
        label="导出DOCX",
        data=export_docx(st.session_state.display_messages),
        file_name=f"对话_{datetime.now(BEIJING_TZ).strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="export_docx_btn"
    )

with st.sidebar:
    st.title("智飞投研")
    st.caption(f"当前时间: {now_ts_display()}")
    st.info(f"当前模型: **{st.session_state.model_name}**")
    st.divider()

    st.subheader("分析方案")
    scheme_cols_1 = st.columns(3)
    scheme_cols_2 = st.columns(3)
    schemes = ["盘前快速分析", "产业链扫描", "卡脖子扫描", "市场行情判断", "资金全景动态", "标的股研报"]
    for i, scheme_name in enumerate(schemes):
        col = scheme_cols_1[i] if i < 3 else scheme_cols_2[i - 3]
        with col:
            if st.button(scheme_name, key=f"scheme_{i}", use_container_width=True):
                st.session_state.scheme = scheme_name
                if is_one_click_scheme(scheme_name):
                    st.session_state._pending_prompt = get_scheme_prompt(scheme_name)
                    st.rerun()
                else:
                    st.toast(f"已切换至: {scheme_name}，请输入标的/事件", icon="ℹ️")
    st.divider()

    st.subheader("快捷工具")
    tool_cols = st.columns(3)
    with tool_cols[0]:
        if st.button("简报", use_container_width=True, key="daily_brief"):
            st.session_state.scheme = "简报"
            st.session_state._pending_prompt = "今日简报"
            st.rerun()
    with tool_cols[1]:
        if st.button("周报", use_container_width=True, key="weekly_report"):
            st.session_state.scheme = "周报"
            st.session_state._pending_prompt = "本周周报"
            st.rerun()
    with tool_cols[2]:
        if st.button("速查", use_container_width=True, key="quick_stock"):
            st.session_state.scheme = "速查"
            st.toast("请输入股票代码或名称，如：300308", icon="ℹ️")
    st.success(f"当前方案: **{st.session_state.scheme}**")
    st.divider()

    st.subheader("Token 监控")
    total_used = st.session_state.get("total_tokens_used", 0)
    BUDGET = 1000000
    usage_ratio = min(total_used / BUDGET, 1.0)
    if usage_ratio >= 0.8:
        st.warning(f"Token使用量已达 {usage_ratio*100:.1f}%")
    elif usage_ratio >= 0.6:
        st.info(f"Token使用量: {usage_ratio*100:.1f}%")
    st.progress(usage_ratio, text=f"{total_used:,} / {BUDGET:,} ({usage_ratio*100:.1f}%)")
    st.caption(f"本会话累计消耗: `{total_used:,}` Tokens")
    st.divider()

    st.subheader("系统状态")
    st.caption("模型自行管理上下文 (RDS)")
    st.divider()

    st.subheader("历史会话")
    history_sessions = get_history_sessions()
    if history_sessions:
        for s in history_sessions:
            sid = s["session_id"]
            title = s["title"]
            rounds = s["rounds"]
            last_time = s["last_time"][:16] if s["last_time"] else ""

            col_a, col_b = st.columns([4, 1])
            with col_a:
                is_current = (sid == st.session_state.session_id)
                label = f"{'🟢 ' if is_current else ''}{title} ({rounds}轮)"
                if st.button(label, key=f"hist_{sid}", use_container_width=True,
                           help=f"最后活跃: {last_time}"):
                    if sid != st.session_state.session_id:
                        msgs = load_session_messages(sid)
                        st.session_state.session_id = sid
                        st.session_state.messages = msgs
                        st.session_state.display_messages = msgs.copy()
                        st.session_state.render_offset = 0
                        st.session_state.history_loaded = True
                        st.rerun()
            with col_b:
                if st.button("🗑", key=f"del_{sid}", help="删除此会话"):
                    if delete_session(sid):
                        if sid == st.session_state.session_id:
                            st.session_state.session_id = str(uuid.uuid4())
                            st.session_state.messages = []
                            st.session_state.display_messages = []
                            st.session_state.render_offset = 0
                            st.session_state.history_loaded = False
                        st.rerun()
    else:
        st.caption("暂无历史会话")
    st.divider()

    if st.button("清空显示", use_container_width=True):
        st.session_state.display_messages = []
        st.session_state.render_offset = 0
        st.rerun()